import json
import requests
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

with open("MsApiTokensBGR.postman_collection.json") as archivo:
    collection = json.load(archivo)

nombre_coleccion = collection['info']['name']
items = collection['item']
print(nombre_coleccion)

def ejecutar_peticion(url, json_data):
    headers = {
        'Content-Type': 'application/json'
    }
    try:
        response = requests.post(url, headers=headers, json=json_data)
        response.raise_for_status()
        return response
    except requests.exceptions.RequestException as e:
        print(f"Error al realizar la petición: {e}")
        return 

def procesar_json_collection():
    for item in items:
        print(item['name'])
        for endpoint in item['item']:
            url = endpoint['request']['url']['raw']
            json_request = json.loads(endpoint['request']['body']['raw'])
            print(type(json_request))
            print(url)
            respuesta = ejecutar_peticion(url, json_request)
            print('res',respuesta)

def generar_documentacion_word(datos):
    try:
        doc = Document()
        tabla = doc.add_table(rows=1, cols=6)
        tabla.style = 'Table Grid'

        encabezados = ["Nombre", "Tipo", "Tamaño", "Requerido", "Descripcion", "Valores de Ejemplo"]
        hdr_cells = tabla.rows[0].cells
        
        for i, encabezado in enumerate(encabezados):
            parrafo = hdr_cells[i].paragraphs[0]
            run = parrafo.add_run(encabezado)
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(10)
              
        # Añadir datos a la tabla
        for key, value in datos.items():
            fila = tabla.add_row().cells
            fila[0].text = key
            fila[1].text = "String"
            fila[2].text = "N/A"
            fila[2].text = "SI"
            fila[5].text = str(value)
        
        doc.save('datos_plantilla.docx')
    except requests.exceptions.RequestException as e:
        print(f"Error al generar el documento: {e}")
        return

datos = {
        "NombrePlantilla": "TOKEN DATOS PLANTILLA - BDUARTE",
        "CodigoProducto": "44",
        "LongitudToken": 6,
        "IncluyeLetras": False,
        "EsCalendarizado": False,
        "CodigoCanal": "0",
        "CodigoMedioInvocacion": "0",
        "PlantillaNotificacion": "0000000000000000",
        "TituloNotificacion": "CÓDIGO CUENTA ZZZZ",
        "DescripcionNotificacion": "Tu código para apertura de Cuenta ZZZZ es:",
        "AsuntoNotificacion": "Código de Seguridad"
    }
generar_documentacion_word(datos)